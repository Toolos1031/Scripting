from docx import Document
import pandas as pd

document = Document()

df = pd.read_excel(r"D:\Fotowoltaika\ORLEN PLOCK\plock_2_pomiary.xlsx")
df = df.fillna(0)

image_path = r"D:\Fotowoltaika\ORLEN PLOCK\Zdjecia_Pole_2\z_power_pointa"


def createTableWithRGBandThermo(Strona, Numer_wskazania, Stol, Panel, Rodzaj_uszkodzenia, Photo):
    table = document.add_table(rows = 4, cols = 5)
    table.allow_autofit = True
    table.style = "TableGrid"

    Strona1 = Strona + 1

    Rgb_photo = f"{image_path}\{Strona}_3.jpeg"
    Thermo_photo = f"{image_path}\{Strona1}_2.jpeg"

    Rgb_photo1 = f"{image_path}\{Strona}_1.jpeg"
    Rgb_photo2 = f"{image_path}\{Strona}_2.jpeg"
    Thermo_photo1 = f"{image_path}\{Strona}_2.jpeg"

    a = table.cell(2, 1)
    b = table.cell(2, 4)
    a.merge(b)

    c = table.cell(3, 1)
    d = table.cell(3, 4)
    c.merge(d)

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Numer wskazania"
    hdr_cells[1].text = "Stół"
    hdr_cells[2].text = "Panel"
    hdr_cells[3].text = "Rodzaj uszkodzenia"
    hdr_cells[4].text = "Numer zdjęcia"

    input_cells = table.rows[1].cells
    input_cells[0].text = Numer_wskazania
    input_cells[1].text = Stol
    input_cells[2].text = Panel
    input_cells[3].text = Rodzaj_uszkodzenia
    input_cells[4].text = Photo

    hdr_rows = table.columns[0].cells
    hdr_rows[2].text = "Zdjęcie RGB"
    hdr_rows[3].text = "Zdjęcie termalne"

    hdr2_rows = table.columns[1].cells
    try:
        paragraph1 = hdr2_rows[2].paragraphs[0]
        run = paragraph1.add_run()
        run.add_picture(Rgb_photo, width = 4200000, height = 2800000)
    except:
        try:
            paragraph1 = hdr2_rows[2].paragraphs[0]
            run = paragraph1.add_run()
            run.add_picture(Rgb_photo1, width = 4200000, height = 2800000)
        except:
            paragraph1 = hdr2_rows[2].paragraphs[0]
            run = paragraph1.add_run()
            run.add_picture(Rgb_photo2, width = 4200000, height = 2800000)

    try:
        paragraph2 = hdr2_rows[3].paragraphs[0]
        run1 = paragraph2.add_run()
        run1.add_picture(Thermo_photo, width = 4200000, height = 2800000)
    except:
        print(f"FAILED AT PAGE {Strona} THERMO")
        #paragraph2 = hdr2_rows[3].paragraphs[0]
        #run = paragraph2.add_run()
        #run.add_picture(Thermo_photo1, width = 4200000, height = 2800000)
    
    document.add_page_break()

def createTableWithMeasurement(Strona, Numer_wskazania, Stol, Panel, Rodzaj_uszkodzenia, Thermo, min_temp, mean_temp, max_temp):
    table = document.add_table(rows = 8, cols = 5)
    table.allow_autofit = True
    table.style = "TableGrid"

    photo = f"{image_path}\{Strona}_2.jpeg"
    photo1 = f"{image_path}\{Strona}_2.jpeg"


    c = table.cell(2, 1)
    d = table.cell(2, 4)
    c.merge(d)

    e = table.cell(3, 0)
    f = table.cell(3, 2)
    e.merge(f)

    g = table.cell(3, 3)
    h = table.cell(3, 4)
    g.merge(h)

    i = table.cell(4, 1)
    j = table.cell(4, 2)
    i.merge(j)

    k = table.cell(5, 1)
    l = table.cell(5, 2)
    k.merge(l)

    m = table.cell(6, 1)
    n = table.cell(6, 2)
    m.merge(n)

    o = table.cell(7, 1)
    p = table.cell(7, 2)
    o.merge(p)

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Numer wskazania"
    hdr_cells[1].text = "Stół"
    hdr_cells[2].text = "Panel"
    hdr_cells[3].text = "Rodzaj uszkodzenia"
    hdr_cells[4].text = "Numer zdjęcia"

    input_cells = table.rows[1].cells
    input_cells[0].text = Numer_wskazania
    input_cells[1].text = Stol
    input_cells[2].text = Panel
    input_cells[3].text = Rodzaj_uszkodzenia
    input_cells[4].text = Thermo

    hdr_rows = table.columns[0].cells
    hdr_rows[2].text = "Zdjęcie termalne z pomiarem"
    hdr_rows[4].text = "Data"
    hdr_rows[5].text = "Minimalna"
    hdr_rows[6].text = "Średnia"
    hdr_rows[7].text = "Maksymalna"

    input_rows = table.columns[1].cells
    input_rows[4].text = "29.08.2024"
    input_rows[5].text = min_temp
    input_rows[6].text = mean_temp
    input_rows[7].text = max_temp

    hdr2_rows = table.columns[3].cells
    hdr2_rows[4].text = "Odległość"
    hdr2_rows[5].text = "Wilgotność"
    hdr2_rows[6].text = "Emisyjność"
    hdr2_rows[7].text = "Temperatura odbita"

    hdr3_rows = table.columns[4].cells
    hdr3_rows[4].text = "22 m"
    hdr3_rows[5].text = "70%"
    hdr3_rows[6].text = "0.90"
    hdr3_rows[7].text = "25.0°C"

    hdr2_cells = table.rows[3].cells
    hdr2_cells[0].text = "POMIAR SQ1"
    hdr2_cells[3].text = "PARAMETRY"

    try:
        hdr3_rows = table.columns[1].cells
        paragraph = hdr3_rows[2].paragraphs[0]
        run = paragraph.add_run()
        run.add_picture(photo, width = 4200000, height = 2800000)
    except:
        hdr3_rows = table.columns[1].cells
        paragraph = hdr3_rows[2].paragraphs[0]
        run = paragraph.add_run()
        run.add_picture(photo1, width = 4200000, height = 2800000)

    document.add_page_break()

for index, row in df.iterrows():
    Strona = row[0] + 1
    Numer_wskazania = row[1]
    Stol = row[2]
    Panel = row[3]
    Rodzaj_uszkodzenia = row[4]
    Thermo = row[6]
    Rgb = row[7]
    min_temp = row[8]
    mean_temp = row[9]
    max_temp = row[10]
    Photo = str(Thermo) + "\n" + str(Rgb)

    

    if Rgb != 0:
        createTableWithRGBandThermo(Strona, str(Numer_wskazania), Stol, Panel, Rodzaj_uszkodzenia, Photo)
    else:
        createTableWithMeasurement(Strona, str(Numer_wskazania), Stol, Panel, Rodzaj_uszkodzenia, Thermo, min_temp, mean_temp, max_temp)

document.save(r"D:\Fotowoltaika\ORLEN PLOCK\plock_2_raport.docx")